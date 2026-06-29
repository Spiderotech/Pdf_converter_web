#!/usr/bin/env python3
import sys
from pathlib import Path

import pdfplumber
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor


MIN_TABLE_ROWS = 2
MIN_TABLE_COLS = 2


def normalize_cell_text(value) -> str:
    if value is None:
        return ""

    return " ".join(str(value).replace("\x00", "").split())


def has_meaningful_table(table_data) -> bool:
    if not table_data or len(table_data) < MIN_TABLE_ROWS:
        return False

    max_columns = max((len(row) for row in table_data if row), default=0)
    filled_cells = sum(
        1
        for row in table_data
        for cell in row
        if normalize_cell_text(cell)
    )

    return max_columns >= MIN_TABLE_COLS and filled_cells >= MIN_TABLE_ROWS * MIN_TABLE_COLS


def add_table(document: Document, table_data) -> None:
    rows = len(table_data)
    columns = max(len(row) for row in table_data)
    table = document.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    for row_index, row_data in enumerate(table_data):
        cells = table.rows[row_index].cells

        for column_index in range(columns):
            text = normalize_cell_text(row_data[column_index] if column_index < len(row_data) else "")
            cell = cells[column_index]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cell.text = text

            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)


def add_text(document: Document, text: str) -> None:
    for line in text.splitlines():
        cleaned_line = " ".join(line.split())

        if not cleaned_line:
            continue

        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(cleaned_line)
        run.font.size = Pt(10)


def add_page_heading(document: Document, page_number: int, total_pages: int) -> None:
    if total_pages <= 1:
        return

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(f"Page {page_number}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(90, 90, 90)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)


def convert_pdf_to_structured_docx(input_pdf: Path, output_docx: Path) -> None:
    document = Document()
    configure_document(document)

    with pdfplumber.open(input_pdf) as pdf:
        total_pages = len(pdf.pages)

        for page_index, page in enumerate(pdf.pages):
            if page_index > 0:
                document.add_section(WD_SECTION.NEW_PAGE)

            add_page_heading(document, page_index + 1, total_pages)

            tables = [
                table
                for table in page.extract_tables()
                if has_meaningful_table(table)
            ]

            if tables:
                for table_index, table_data in enumerate(tables):
                    if table_index > 0:
                        document.add_paragraph()

                    add_table(document, table_data)
            else:
                text = page.extract_text(layout=False, x_tolerance=2, y_tolerance=4) or ""
                add_text(document, text)

            if not tables and not (page.extract_text() or "").strip():
                paragraph = document.add_paragraph()
                run = paragraph.add_run(
                    "No selectable text was found on this page. This experimental converter "
                    "does not OCR scanned/image-only PDFs."
                )
                run.italic = True
                run.font.size = Pt(10)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_docx)


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: pdf_to_structured_docx.py <input.pdf> <output.docx>", file=sys.stderr)
        return 2

    input_pdf = Path(sys.argv[1]).resolve()
    output_docx = Path(sys.argv[2]).resolve()

    if not input_pdf.exists():
        print(f"Input PDF not found: {input_pdf}", file=sys.stderr)
        return 1

    convert_pdf_to_structured_docx(input_pdf, output_docx)

    if not output_docx.exists() or output_docx.stat().st_size == 0:
        print("DOCX output was not created.", file=sys.stderr)
        return 1

    print(f"Structured DOCX created: {output_docx}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
